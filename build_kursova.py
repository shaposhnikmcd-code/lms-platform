# -*- coding: utf-8 -*-
"""Build the kursova .docx from kursova_pastor.md with strict formatting."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\Shapo\Downloads\kursova_pastor.md"
OUT = r"C:\Users\Shapo\Downloads\БИБЛЕЙСКИЙ ПОРТРЕТ ПАСТОРА.docx"

with open(SRC, encoding="utf-8") as f:
    md = f.read()

# Cut out the title page and biblio sections by markers
# We'll manually rebuild title; content goes from "**СОДЕРЖАНИЕ**" forward.
content_start = md.index("**СОДЕРЖАНИЕ**")
md_content = md[content_start:]

doc = Document()

# ---- Page setup ----
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

# ---- Default style ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(14)
# East Asian font for Cyrillic
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.insert(0, rfonts)
rfonts.set(qn("w:ascii"), "Times New Roman")
rfonts.set(qn("w:hAnsi"), "Times New Roman")
rfonts.set(qn("w:cs"), "Times New Roman")
rfonts.set(qn("w:eastAsia"), "Times New Roman")

pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Cm(1.25)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ---- Page numbers top center, none on first page ----
section.different_first_page_header_footer = True

def add_page_number(hdr):
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    run.font.name = "Times New Roman"; run.font.size = Pt(14)

add_page_number(section.header)
# first page header empty (no number on titul)

# ---- Helpers ----
def add_p(text, *, bold=False, italic=False, align="justify", indent=True,
          size=14, space_before=0, space_after=0, center=False):
    p = doc.add_paragraph()
    aligns = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
              "center": WD_ALIGN_PARAGRAPH.CENTER,
              "left": WD_ALIGN_PARAGRAPH.LEFT,
              "right": WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = aligns["center" if center else align]
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p

def add_rich(text, **kw):
    """Render text with **bold** segments."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
    return p

def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def heading_main(text):
    """Top-level heading: UPPERCASE bold center, новая страница."""
    page_break()
    add_p(text, bold=True, center=True, indent=False, space_after=18)

def heading_sub(text):
    """Subheading 1.1 — abzac, lowercase bold (already lowercase)."""
    add_p(text, bold=True, indent=True, space_before=6, space_after=6, align="left")

# ---- TITLE PAGE in a frame (single-cell table with borders) ----
titul = doc.add_table(rows=1, cols=1)
titul.autofit = False
titul.columns[0].width = Cm(16.5)
tc = titul.cell(0, 0)
tc.width = Cm(16.5)
tcPr = tc._tc.get_or_add_tcPr()
tcBorders = OxmlElement("w:tcBorders")
for side in ("top", "left", "bottom", "right"):
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "18")
    b.set(qn("w:color"), "000000")
    tcBorders.append(b)
tcPr.append(tcBorders)
# remove cell shading (no background)
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"), "FFFFFF")
tcPr.append(shd)

def tcell(text, *, bold=False, size=14, align="center", first=False):
    if first:
        p = tc.paragraphs[0]
    else:
        p = tc.add_paragraph()
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
    return p

tcell("МЕЖДУНАРОДНЫЙ УНИВЕРСИТЕТ «ВИДЕНИЕ»", bold=True, first=True)
tcell("Филиал № 1", bold=True)
tcell("")
tcell("Магистратура")
for _ in range(4): tcell("")
tcell("КУРСОВАЯ РАБОТА", bold=True, size=18)
tcell("")
tcell("по предмету: «Пастырское служение»")
tcell("")
tcell("на тему:")
tcell("«БИБЛЕЙСКИЙ ПОРТРЕТ ПАСТОРА»", bold=True, size=16)
for _ in range(6): tcell("")
tcell("Выполнила: Шапошник Татьяна Васильевна", align="right")
tcell("Год обучения: 1 квартал 2025 — 4 квартал 2026", align="right")
tcell("Преподаватель: А. А. Руднев", align="right")
tcell("Проверяющий: ______________________", align="right")
for _ in range(5): tcell("")
tcell("Киев, апрель 2026")

page_break()

# ---- СОДЕРЖАНИЕ ----
add_p("СОДЕРЖАНИЕ", bold=True, center=True, indent=False, space_after=18)
toc_items = [
    ("ВВЕДЕНИЕ", "3"),
    ("1. БИБЛЕЙСКИЕ ОСНОВАНИЯ ПАСТЫРСКОГО СЛУЖЕНИЯ", "5"),
    ("    1.1. Понятие пастыря в Ветхом и Новом Завете", "5"),
    ("    1.2. Призвание и помазание как основа пасторского служения", "9"),
    ("2. ХАРАКТЕР И КАЧЕСТВА ПАСТОРА ПО ПОСЛАНИЯМ АПОСТОЛА ПАВЛА", "13"),
    ("    2.1. Требования к епископу в 1 Тим. 3 и Тит. 1: каким должен быть пастор", "13"),
    ("    2.2. Семья, духовная и нравственная зрелость пастора", "17"),
    ("3. ПРАКТИЧЕСКИЕ АСПЕКТЫ ПАСТЫРСКОГО СЛУЖЕНИЯ", "21"),
    ("    3.1. Каким Бог хочет видеть пастора: подражание, пример и влияние", "21"),
    ("    3.2. Слабые места пасторского служения и пути их преодоления", "25"),
    ("ЗАКЛЮЧЕНИЕ", "29"),
    ("БИБЛИОГРАФИЯ", "31"),
]
for title, pg in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    # tab leader
    tab_stops = p.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(title); r.font.name = "Times New Roman"; r.font.size = Pt(14)
    r2 = p.add_run("\t" + pg); r2.font.name = "Times New Roman"; r2.font.size = Pt(14)

# ---- Parse the content from .md ----
# Strip the manually-built sections (содержание already done above) — start after СОДЕРЖАНИЕ block
rest = md_content
# remove the soderzhanie block we already rendered
rest = re.sub(r"\*\*СОДЕРЖАНИЕ\*\*.*?(?=---)", "", rest, count=1, flags=re.S)

# Remove horizontal rules
rest = rest.replace("---", "")

# Process line by line / paragraph by paragraph
paragraphs = [p.strip() for p in rest.split("\n\n")]

H1_RE = re.compile(r"^\*\*(\d+)\.\s+(.+)\*\*$")
H2_RE = re.compile(r"^\*\*(\d+\.\d+)\.\s+(.+)\*\*$")
PLAIN_BOLD_RE = re.compile(r"^\*\*([^*].+)\*\*$")  # ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ etc

for para in paragraphs:
    if not para:
        continue
    if para.startswith("**БИБЛИОГРАФИЯ**"):
        break  # bibliography handled manually
    m1 = H1_RE.match(para)
    m2 = H2_RE.match(para)
    mp = PLAIN_BOLD_RE.match(para)
    if m1:
        heading_main(f"{m1.group(1)}. {m1.group(2).upper()}")
    elif m2:
        heading_sub(f"{m2.group(1)}. {m2.group(2)}")
    elif mp and mp.group(1) in ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"):
        heading_main(mp.group(1))
    elif para.startswith("1. ") or para.startswith("2. ") or para.startswith("3. ") or \
         para.startswith("4. ") or para.startswith("5. ") or para.startswith("6. "):
        # numbered list — multi-line
        for line in para.split("\n"):
            line = line.strip()
            if not line: continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(line); r.font.name = "Times New Roman"; r.font.size = Pt(14)
    else:
        # plain paragraph (may have inline bold)
        # collapse newlines within paragraph
        text = " ".join(para.split())
        add_rich(text)

# ---- БИБЛИОГРАФИЯ ----
heading_main("БИБЛИОГРАФИЯ")
biblio = [
    "Адамс Дж. Э. Учебник по церковной дисциплине / Пер. с англ. — Одесса: Богомыслие, 2003. — 152 с.",
    "Алфорд Х., Нотон М. Менеджмент, когда вера имеет значение: Христианские социальные принципы в управлении современной организацией / Пер. с англ. — Москва: Дух и литера, 2007. — 384 с.",
    "Бланшар К., Ходжес Ф. Руководи, как Иисус / Пер. с англ. — Москва: Эксмо, 2008. — 256 с.",
    "Гарднер Л. Утверждение истины: Осмысленный взгляд на основания христианства / Пер. с англ. — Москва: Триада, 2005. — 272 с.",
    "Гец Дж. Внимательный взгляд на церковь / Пер. с англ. — BEE International, 1990. — 318 с.",
    "Гетц Дж. А. Мера мужского характера / Пер. с англ. — Санкт-Петербург: Библия для всех, 2001. — 224 с.",
    "Даффилд Г. П., Ван Клив Н. М. Основы пятидесятнического богословия / Пер. с англ. — Москва: Жизнь, 2007. — 656 с.",
    "Джойнер Р. Лидер, менеджер и пять принципов успеха / Пер. с англ. — Москва: МираТ, 2003. — 192 с.",
    "Костюкова Т. А., Петрова Г. И. Христианская педагогика в современном образовательном пространстве. — Томск: Издательство Томского государственного педагогического университета, 2001. — 168 с.",
    "Максвелл Дж. 21 обязательное качество лидера. — Минск: Попурри, 2002. — 176 с.",
    "Рассел Б. Когда Бог созидает церковь: Десять принципов растущей церкви / Пер. с англ. — Санкт-Петербург: Шандал, 2004. — 208 с.",
    "Сазерлэнд Д., Новери К. 33 закона управления: Принципы плодотворной жизни / Пер. с англ. — Санкт-Петербург: Шандал, 2005. — 192 с.",
    "Трит К. Лидерство, ведущее к росту церкви. Серия «Для лидеров и пасторов» / Пер. с англ. — Москва: Триада, 2003. — 144 с.",
    "Уэрсби У. Быть Божьим слугой / Пер. с англ. — Санкт-Петербург: Библия для всех, 2002. — 192 с.",
    "Хеґберг Дж. О., Ґюліч Р. А. Критична подорож: Етапи життя віри / Пер. з англ. — Київ: УБТС, 2010. — 256 с.",
    "Хэррис Р. Полное наглядное пособие по терапии принятия и ответственности / Пер. с англ. — Москва: Бомбора, 2021. — 496 с.",
    "Ялом И. Дар психотерапии / Пер. с англ. — Москва: Эксмо, 2009. — 352 с.",
]
for i, item in enumerate(biblio, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"{i}. {item}")
    r.font.name = "Times New Roman"; r.font.size = Pt(14)

# ---- Remove any document background ----
settings = doc.settings.element
# remove w:background if any
bg = doc.element.find(qn("w:background"))
if bg is not None:
    doc.element.remove(bg)

doc.save(OUT)
print(f"Saved: {OUT}")
